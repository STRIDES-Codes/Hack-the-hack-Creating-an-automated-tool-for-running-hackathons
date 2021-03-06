from docx import Document
from docx.shared import Inches
import json
import requests
import pandas
import os, sys

def createDoc():
    doca = Document()
    instructions = '[This template contains all of the necessary sections for completing a document for the F1000R Hackathons channel.  Do not remove any sections. Do not change the order of any sections.  Figures should be submitted as separate files.  Detailed author instructions are available at http://f1000research.com/for-authors/article-guidelines/software-tool-articles.]'
    firstTitle = 'Title (sentence-case, not title case)'
    doca.add_paragraph().add_run(instructions).italic = True
    
    doca.add_heading(firstTitle, 0)
    word1 = 'Authors'    
    doca.add_paragraph().add_run(word1).bold = True
    p2 = doca.add_paragraph('Author Name 1')
    p2.add_run('\nAuthor email address')
    p2.add_run('\nAuthor affiliation, including full address with zip code')
    p3 = doca.add_paragraph('Author Name 2')
    p3.add_run('\nAuthor email address')
    p3.add_run('\nAuthor affiliation, including full address with zip code')
    doca.add_heading('Abstract', 1)
    doca.add_paragraph('Insert abstract here (up to 300 words, no citations, spell out all abbreviations).')
    doca.add_heading('Keywords', 1)
    doca.add_paragraph('Provide up to 8 keywords, comma-separated')
    doca.add_heading('Introduction', 1)
    doca.add_paragraph('Insert introduction here. This section should provide context as to why the software tool was developed and what need it addresses')
    doca.add_heading('Methods', 1)
    doca.add_paragraph('Insert methods here, with the two required subsections as described below:')
    doca.add_heading('Implementation', 1)
    doca.add_paragraph('This section describes how the tool works and relevant technical details for implementation.')
    doca.add_heading('Operation', 1)
    doca.add_paragraph('This section should include the minimal system requirements needed to run the software and an overview of the workflow')
    doca.add_heading('Results', 1)
    doca.add_paragraph('Include if the paper includes novel data or analyses; should be written as a traditional results section (otherwise, include a Use Cases section).')
    doca.add_heading('Use Cases', 1)
    doca.add_paragraph('If not including a Results section, include this section. Examples of input and output files should be provided with some explanatory context. Any novel or complex variable parameters should be explained in sufficient detail to enable users to understand and use the tools functionality.')
    doca.add_heading('Conclusion and next steps', 1)
    doca.add_paragraph('This section should include a brief discussion of allowances made (if any) for controlling bias or unwanted sources of variability, and the limitations of any novel datasets. Also include any next steps for future development (whether your group actually plans to do this or these steps are just included a guidance for potential future development).')
    doca.add_heading('Data and software availability', 1)
    doca.add_paragraph('Source code for new software must be made openly, and permanently available in a structured repository such as Zenodo (the F1000Research team can assist with deposition), as well as uploaded to a Version Control System (VCS) such as GitHub, BitBucket or SourceForge. Please provide details in a section entitled ???Software availability???, listing the repository and the license under which the software can be used in the article. Source code must be assigned an open license.')
    doca.add_heading('Suggested Reviewers', 1)
    doca.add_paragraph('Please pick ten suggested reviewers with whom you have not published in the last three years and who do not work in the same department.  Papers can not be sent to F1000 research without such a list.')
    doca.add_heading('Author contributions', 1)
    doca.add_paragraph('F1000R uses the CRediT Taxonomy for author contributions.  For each author, list their contribution(s) from the list below.  Anyone who contributed in another capacity or otherwise does not meet the criteria for authorship (e.g. they did not review the final manuscript) should be included in the acknowledgements.')
    items = [
        ['Conceptualization', 'Ideas; formulation or evolution of overarching research goals and aims.'], 
        ['Data Curation', 'Management activities to annotate (produce metadata), scrub data and maintain research data (including software code, where it is necessary for interpreting the data itself) for initial use and later reuse.'],
        ['Formal Analysis', 'Application of statistical, mathematical, computational, or other formal techniques to analyze or synthesize study data.'], 
        ['Funding Aquisition', 'Acquisition of the financial support for the project leading to this publication.'],
        ['Investigation', 'Conducting a research and investigation process, specifically performing the experiments, or data/evidence collection.'],
        ['Methodology', 'Development or design of methodology; creation of models.'],
        ['Project Administration', 'Management and coordination responsibility for the research activity planning and execution.'],
        ['Resources', 'Provision of study materials, reagents, materials, patients, laboratory samples, animals, instrumentation, computing resources, or other analysis tools.'],
        ['Software', 'Programming, software development; designing computer programs; implementation of the computer code and supporting algorithms; testing of existing code components.'],
        ['Supervision', 'Oversight and leadership responsibility for the research activity planning and execution, including mentorship external to the core team.'],
        ['Validation', 'Verification, whether as a part of the activity or separate, of the overall replication/reproducibility of results/experiments and other research outputs.'],
        ['Visualization', 'Preparation, creation and/or presentation of the published work, specifically visualization/data presentation.'],
        ['Writing - Original Draft Preparation', 'Creation and/or presentation of the published work, specifically writing the initial draft (including substantive translation).'],
        ['Writing - Review and Editing', 'Preparation, creation and/or presentation of the published work by those from the original research group, specifically critical review, commentary or revision ??? including pre- or post-publication stages.']
            ]
    contribtable = doca.add_table(rows = 1, cols = 2)
    contribtable.style = 'Table Grid'
    hdr_cells = contribtable.rows[0].cells
    run = hdr_cells[0].paragraphs[0].add_run('Contributor Role')
    run.bold = True
    run = hdr_cells[1].paragraphs[0].add_run('Role Definition')
    run.bold = True
#    hdr_cells[0].text = 'Contributor Role'
#   hdr_cells[1].text = 'Role Definition'
    
    for contributor, role in items:
        row_Cells = contribtable.add_row().cells
        row_Cells[0].text = str(contributor)
        row_Cells[1].text = str(role)
    doca.add_heading('Competing interests', 1)
    doca.add_paragraph('Note any financial, personal, or professional competing interests for any of the authors that could be construed to unduly influence the content of the article. If none, include the text ???No competing interests were disclosed.???')
    doca.add_heading('Grant information', 1)
    doca.add_paragraph('Include funding if relevant (including funding from authors??? employers if relevant, and any relevant grant funding).  If none, include the text ???The author(s) declared that no grants were involved in supporting this work???.')
    doca.add_heading('Acknowledgements', 1)
    doca.add_paragraph('This section should acknowledge anyone who contributed to the research or the writing of the article but who does not qualify as an author; please clearly state how they contributed. Authors should obtain permission to include the name and affiliation, from all those mentioned in the Acknowledgments section.')
    doca.add_heading('References', 1)
    refs = doca.add_paragraph('Instructions on using the F1000R Google docs plug in for reference management: http://f1000.com/work/faq/google-docs-add-on/1')
    refs.add_run('Instructions on using the F1000R Word plug in for reference management: http://f1000.com/work/faq/word-plugin')
    doca.add_heading('Figures and Tables', 1)
    doca.add_paragraph('All figures and tables should be cited and discussed in the article text. Figure legends and tables should be added at the end of the manuscript. Tables should be formatted using the ???insert table??? function in Word, or provided as an Excel file. Files for figures should be uploaded as separate files through the submission system. Each figure or table should have a concise title of no more than 15 words. A legend for each figure and table should also be provided that briefly describes the key points and explains any symbols and abbreviations used. The legend should be sufficiently detailed so that the figure or table can stand alone from the main text. ')
    doca.save('template.docx')
    
def uploadDoc(auth_key, projectName):
    
    headers = {"Authorization": "Bearer " + auth_key}
    para = {
        "name": projectName,
    }
    files = {
        'data': ('metadata', json.dumps(para), 'application/json; charset=UTF-8'),
        'file': open("./template.docx", "rb")
    }
    r = requests.post(
        "https://www.googleapis.com/upload/drive/v3/files?uploadType=multipart",
        headers=headers,
        files=files
    )
    print(r.text,file=sys.stderr)

def uploadDocs(project_dictionary):
    allProjects = project_dictionary['teams']
    for project in allProjects:
            uploadDoc(project_dictionary['host']['drive_auth_key'], project['team_name'])
            
def closeDoc():
    os.remove("template.docx")
    print("File Removed!",file=sys.stderr)
            
#createDoc()
#uploadDocs(test_object)
    
